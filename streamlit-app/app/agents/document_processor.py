from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Optional
from pathlib import Path
from langchain_core.documents import Document
import streamlit as st
import PyPDF2
import docx
import io
from langchain_community.llms.bedrock import Bedrock

class DocumentProcessor:
    def __init__(self, llm: Bedrock, docs_path: Optional[Path] = None):
        """Initialize DocumentProcessor with Bedrock LLM."""
        self.llm = llm
        self.docs_path = docs_path
        # Adjust chunk size for better compatibility with Bedrock models
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Smaller chunks for better embedding
            chunk_overlap=100,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ";"]
        )

    def process_documents(self, files, is_uploaded: bool = False) -> List[Document]:
        """Process either uploaded files or files from disk."""
        documents = []
        
        if is_uploaded:
            # Process uploaded files
            for file in files:
                try:
                    content = ""
                    file_extension = Path(file.name).suffix.lower()
                    
                    if file_extension == '.pdf':
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                        content = "\n".join(page.extract_text() for page in pdf_reader.pages)
                    
                    elif file_extension == '.txt':
                        content = file.read().decode('utf-8')
                    
                    elif file_extension == '.docx':
                        doc = docx.Document(io.BytesIO(file.read()))
                        content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    
                    if content.strip():
                        metadata = {
                            "source": file.name,
                            "type": file_extension[1:],
                            "size": file.size
                        }
                        texts = self.text_splitter.split_text(content)
                        documents.extend([
                            Document(page_content=text, metadata=metadata)
                            for text in texts
                        ])
                
                except Exception as e:
                    st.warning(f"Error processing file {file.name}: {str(e)}")
                    continue
        else:
            # Process files from disk
            if not self.docs_path or not self.docs_path.exists():
                return []
                
            for file_path in self.docs_path.glob("*.*"):
                try:
                    content = ""
                    if file_path.suffix.lower() == '.pdf':
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            content = "\n".join(page.extract_text() for page in pdf_reader.pages)
                    
                    elif file_path.suffix.lower() == '.txt':
                        with open(file_path, 'r', encoding='utf-8') as file:
                            content = file.read()
                    
                    elif file_path.suffix.lower() == '.docx':
                        doc = docx.Document(file_path)
                        content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    
                    if content.strip():
                        metadata = {
                            "source": file_path.name,
                            "type": file_path.suffix[1:],
                            "size": file_path.stat().st_size
                        }
                        texts = self.text_splitter.split_text(content)
                        documents.extend([
                            Document(page_content=text, metadata=metadata)
                            for text in texts
                        ])
                
                except Exception as e:
                    print(f"Error processing file {file_path}: {str(e)}")
                    continue
        
        return documents

    def load_and_split(self) -> List[Document]:
        """Load documents from the specified path and split them into chunks."""
        return self.process_documents(None, is_uploaded=False)

    def process_uploaded_files(self, uploaded_files: List[st.runtime.uploaded_file_manager.UploadedFile]) -> List[Document]:
        """Process uploaded files and convert them to Document objects."""
        return self.process_documents(uploaded_files, is_uploaded=True) 